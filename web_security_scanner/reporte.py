import json
from datetime import datetime
import html
import openpyxl
from docx import Document
from flask import Flask, send_file, request
from fpdf import FPDF
import re


app = Flask(__name__)

def generar_reporte_json(resultados, nombre_archivo="scan_results.json"):
    """Genera un reporte en formato JSON"""
    try:
        with open(nombre_archivo, "w", encoding="utf-8") as f:
            json.dump({
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "resultados": resultados
            }, f, indent=4, ensure_ascii=False)
        print(f"[+] Reporte JSON generado: {nombre_archivo}")
    except Exception as e:
        print(f"[!] Error al generar el reporte JSON: {e}")

def generar_reporte_html(resultados, nombre_archivo="scan_results.html"):
    """Genera un reporte en formato HTML"""
    try:
        html_content = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Reporte Web Security Scanner</title>
    <style>
        body {{ font-family: Arial, sans-serif; background: #f4f4f4; color: #222; }}
        h1 {{ background: #0078d7; color: #fff; padding: 10px; }}
        .section {{ background: #fff; margin: 20px 0; padding: 15px; border-radius: 8px; box-shadow: 0 2px 6px #ccc; }}
        .vuln {{ color: #c00; font-weight: bold; }}
        .safe {{ color: #080; font-weight: bold; }}
        .info {{ color: #0078d7; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
        th, td {{ border: 1px solid #ccc; padding: 6px 10px; text-align: left; }}
        th {{ background: #eee; }}
    </style>
</head>
<body>
    <h1>Reporte Web Security Scanner</h1>
    <div class="section info">
        <b>Descargar reportes:</b>
        <ul>
            <li><a href="/descargar/excel">Descargar Excel (.xlsx)</a></li>
            <li><a href="/descargar/word">Descargar Word (.docx)</a></li>
            <li><a href="/descargar/pdf">Descargar PDF (.pdf)</a></li>
        </ul>
        <b>Fecha:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
    </div>
    <div class="section">
        <h2>Información del servidor</h2>
        <ul>
            <li><b>Server:</b> {resultados.get('server_info', {}).get('server', 'No detectado')}</li>
            <li><b>X-Powered-By:</b> {resultados.get('server_info', {}).get('x_powered_by', 'No detectado')}</li>
            <li><b>Content-Type:</b> {resultados.get('server_info', {}).get('content_type', 'No detectado')}</li>
        </ul>
    </div>
    <div class="section">
        <h2>Vulnerabilidades Detectadas</h2>
        <h3>Inyección SQL</h3>
        {generar_tabla_vulns(resultados.get('sql_injection', []), ["url", "method", "payload", "parameter", "severity"])}
        <h3>XSS</h3>
        {generar_tabla_vulns(resultados.get('xss', []), ["url", "method", "payload", "parameter"])}
        <h3>Inyección NoSQL</h3>
        {generar_tabla_vulns(resultados.get('nosql_injection', []), ["url", "method", "payload", "parameter"])}
        <h3>Redirección Abierta</h3>
        {generar_tabla_vulns(resultados.get('open_redirect', []), ["url", "parameter", "redirected_to"])}
    </div>
    <div class="section">
        <h2>Tecnologías Detectadas</h2>
        <ul>
            {"".join(f"<li><b>{k}:</b> {v}</li>" for k, v in resultados.get('technologies', {}).items())}
        </ul>
    </div>
    <div class="section">
        <h2>Resumen</h2>
        <ul>
            <li><b>Formularios encontrados:</b> {resultados.get('forms_found', 0)}</li>
            <li><b>Parámetros encontrados:</b> {", ".join(resultados.get('parameters_found', []))}</li>
        </ul>
    </div>
    <div class="section">
        <h2>Subdirectorios Encontrados</h2>
        <ul>
            {"".join(f"<li>{html.escape(subdir)}</li>" for subdir in resultados.get('subdirectories', []))}
        </ul>
    </div>
    <div class="section">
        <h2>Subdominios Encontrados</h2>
        <ul>
            {"".join(f"<li>{html.escape(subd)}</li>" for subd in resultados.get('subdomains', []))}
        </ul>
    </div>
    <div class="section info">
        <b>Nota:</b> Este reporte es generado automáticamente por Web Security Scanner.<br>
        Uso educativo y legal únicamente.
    </div>
</body>
</html>
"""
        with open(nombre_archivo, "w", encoding="utf-8") as f:
            f.write(html_content)
        print(f"[+] Reporte HTML generado: {nombre_archivo}")
    except Exception as e:
        print(f"[!] Error al generar el reporte HTML: {e}")

def generar_tabla_vulns(vulns, campos):
    if not vulns:
        return '<span class="safe">No se encontraron vulnerabilidades.</span>'
    # Añade la columna Gravedad
    html_table = '<table><tr>' + ''.join(f"<th>{campo.capitalize()}</th>" for campo in campos) + "<th>Gravedad</th></tr>"
    for vuln in vulns:
        html_table += "<tr>" + "".join(
            f"<td>{html.escape(str(vuln.get(campo, '')))}</td>" for campo in campos
        ) + f"<td>{html.escape(str(vuln.get('severity', 'Desconocida')))}</td></tr>"
    html_table += "</table>"
    return html_table

def generar_reporte_excel(resultados, nombre_archivo="scan_results.xlsx"):
    """Genera un reporte en formato Excel (XLSX)"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Vulnerabilidades"

    # Encabezados
    headers = ["Tipo", "URL", "Método", "Payload", "Parámetro", "Redirección a"]
    ws.append(headers)

    # Helper para agregar filas
    def agregar_vulns(tipo, vulns, campos):
        for vuln in vulns:
            fila = [tipo]
            for campo in ["url", "method", "payload", "parameter", "redirected_to"]:
                fila.append(vuln.get(campo, ""))
            ws.append(fila)

    # Agregar todas las vulnerabilidades
    agregar_vulns("SQL Injection", resultados.get('sql_injection', []), headers)
    agregar_vulns("XSS", resultados.get('xss', []), headers)
    agregar_vulns("NoSQL Injection", resultados.get('nosql_injection', []), headers)
    agregar_vulns("Open Redirect", resultados.get('open_redirect', []), headers)

    wb.save(nombre_archivo)
    print(f"[+] Reporte Excel generado: {nombre_archivo}")

def generar_reporte_word(resultados, nombre_archivo="scan_results.docx"):
    """Genera un reporte en formato Word (DOCX)"""
    doc = Document()
    doc.add_heading('Reporte Web Security Scanner', 0)

    def agregar_vulns(tipo, vulns):
        doc.add_heading(tipo, level=1)
        if not vulns:
            doc.add_paragraph("No se encontraron vulnerabilidades.")
            return
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'URL'
        hdr_cells[1].text = 'Método'
        hdr_cells[2].text = 'Payload'
        hdr_cells[3].text = 'Parámetro'
        hdr_cells[4].text = 'Redirección a'
        for vuln in vulns:
            row_cells = table.add_row().cells
            row_cells[0].text = str(vuln.get('url', ''))
            row_cells[1].text = str(vuln.get('method', ''))
            row_cells[2].text = str(vuln.get('payload', ''))
            row_cells[3].text = str(vuln.get('parameter', ''))
            row_cells[4].text = str(vuln.get('redirected_to', ''))
    
    agregar_vulns("SQL Injection", resultados.get('sql_injection', []))
    agregar_vulns("XSS", resultados.get('xss', []))
    agregar_vulns("NoSQL Injection", resultados.get('nosql_injection', []))
    agregar_vulns("Open Redirect", resultados.get('open_redirect', []))

    doc.save(nombre_archivo)
    print(f"[+] Reporte Word generado: {nombre_archivo}")

def generar_reporte_pdf(resultados=None, pdf_file="scan_results.pdf"):
    """
    Genera un PDF con tablas para cada tipo de vulnerabilidad, similar al HTML.
    Si resultados es None, los carga desde scan_results.json.
    """
    try:
        if resultados is None:
            with open("scan_results.json", encoding="utf-8") as f:
                resultados = json.load(f).get("resultados", {})

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "Reporte Web Security Scanner", ln=True, align="C")
        pdf.set_font("Arial", "", 10)
        pdf.cell(0, 8, f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.ln(4)

        # Información del servidor
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Información del servidor", ln=True)
        pdf.set_font("Arial", "", 10)
        server_info = resultados.get('server_info', {})
        pdf.cell(0, 6, f"Server: {server_info.get('server', 'No detectado')}", ln=True)
        pdf.cell(0, 6, f"X-Powered-By: {server_info.get('x_powered_by', 'No detectado')}", ln=True)
        pdf.cell(0, 6, f"Content-Type: {server_info.get('content_type', 'No detectado')}", ln=True)
        pdf.ln(4)

        # Helper para tablas
        def tabla_vulns(titulo, vulns, campos):
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, titulo, ln=True)
            pdf.set_font("Arial", "B", 9)
            col_widths = [40, 20, 40, 30, 30]
            # Encabezados
            for i, campo in enumerate(campos):
                pdf.cell(col_widths[i], 7, campo.capitalize(), border=1)
            pdf.ln()
            pdf.set_font("Arial", "", 8)
            if not vulns:
                pdf.cell(sum(col_widths), 7, "No se encontraron vulnerabilidades.", border=1, ln=True)
            else:
                for vuln in vulns:
                    for i, campo in enumerate(campos):
                        valor = str(vuln.get(campo, ""))
                        # Trunca si es muy largo
                        if len(valor) > 40:
                            valor = valor[:37] + "..."
                        pdf.cell(col_widths[i], 7, valor, border=1)
                    pdf.ln()
            pdf.ln(2)

        # Tablas de vulnerabilidades
        tabla_vulns("Inyección SQL", resultados.get('sql_injection', []), ["url", "method", "payload", "parameter", "severity"])
        tabla_vulns("XSS", resultados.get('xss', []), ["url", "method", "payload", "parameter", "severity"])
        tabla_vulns("Inyección NoSQL", resultados.get('nosql_injection', []), ["url", "method", "payload", "parameter", "severity"])
        tabla_vulns("Redirección Abierta", resultados.get('open_redirect', []), ["url", "parameter", "redirected_to", "", ""])

        # Tecnologías detectadas
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Tecnologías Detectadas", ln=True)
        pdf.set_font("Arial", "", 10)
        techs = resultados.get('technologies', {})
        if techs:
            for k, v in techs.items():
                pdf.cell(0, 6, f"{k}: {', '.join(v) if isinstance(v, list) else v}", ln=True)
        else:
            pdf.cell(0, 6, "No se detectaron tecnologías.", ln=True)
        pdf.ln(2)

        # Resumen
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Resumen", ln=True)
        pdf.set_font("Arial", "", 10)
        pdf.cell(0, 6, f"Formularios encontrados: {resultados.get('forms_found', 0)}", ln=True)
        pdf.cell(0, 6, f"Parámetros encontrados: {', '.join(resultados.get('parameters_found', []))}", ln=True)
        pdf.ln(2)

        # Subdirectorios y subdominios
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Subdirectorios Encontrados", ln=True)
        pdf.set_font("Arial", "", 10)
        for subdir in resultados.get('subdirectories', []):
            pdf.cell(0, 6, subdir, ln=True)
        pdf.ln(2)
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Subdominios Encontrados", ln=True)
        pdf.set_font("Arial", "", 10)
        for subd in resultados.get('subdomains', []):
            pdf.cell(0, 6, subd, ln=True)

        pdf.output(pdf_file)
        print(f"[+] Reporte PDF generado: {pdf_file}")
    except Exception as e:
        print(f"[!] Error al generar el PDF: {e}")

def get_severity_and_desc(vuln_type):
        """Devuelve la gravedad y una breve descripción según el tipo de vulnerabilidad"""
        if vuln_type == 'sql':
            return (
                "Alta",
                "La inyección SQL permite a un atacante manipular consultas a la base de datos. Referencia: https://owasp.org/www-community/attacks/SQL_Injection"
            )
        elif vuln_type == 'xss':
            return (
                "Alta",
                "El Cross-Site Scripting (XSS) permite ejecutar scripts maliciosos en el navegador de la víctima. Referencia: https://owasp.org/www-community/attacks/xss/"
            )
        elif vuln_type == 'nosql':
            return (
                "Alta",
                "La inyección NoSQL permite manipular consultas en bases de datos NoSQL. Referencia: https://owasp.org/www-community/attacks/NoSQL_Injection"
            )
        elif vuln_type == 'open_redirect':
            return (
                "Media",
                "La redirección abierta puede ser usada para phishing o robo de credenciales. Referencia: https://owasp.org/www-community/attacks/Unvalidated_Redirects_and_Forwards_Cheat_Sheet"
            )
        else:
            return ("Desconocida", "No se pudo determinar la gravedad.")

@app.route('/descargar/<tipo>')
def descargar(tipo):
    # Carga los resultados desde el HTML generado (parseando el JSON embebido o leyendo el archivo JSON)
    # Aquí usamos el archivo scan_results.json que se genera junto con el HTML
    try:
        with open('scan_results.json', encoding='utf-8') as f:
            resultados = json.load(f)['resultados']
    except Exception as e:
        return f"No se pudo cargar la información del reporte: {e}", 500

    if tipo == 'excel':
        generar_reporte_excel(resultados)
        return send_file('scan_results.xlsx', as_attachment=True)
    elif tipo == 'word':
        generar_reporte_word(resultados)
        return send_file('scan_results.docx', as_attachment=True)
    elif tipo == 'pdf':
        generar_reporte_pdf(resultados)
        return send_file('scan_results.pdf', as_attachment=True)
    else:
        return "Tipo de reporte no soportado", 400

# Solo genera Excel y Word cuando el usuario lo descarga desde el HTML
@app.route('/descargar/excel')
def descargar_excel():
    with open('scan_results.json', encoding='utf-8') as f:
        resultados = json.load(f)['resultados']
    generar_reporte_excel(resultados)
    return send_file('scan_results.xlsx', as_attachment=True)

@app.route('/descargar/word')
def descargar_word():
    with open('scan_results.json', encoding='utf-8') as f:
        resultados = json.load(f)['resultados']
    generar_reporte_word(resultados)
    return send_file('scan_results.docx', as_attachment=True)

@app.route('/descargar/pdf')
def descargar_pdf():
    # Genera el PDF solo cuando se solicita
    generar_reporte_pdf()
    return send_file('scan_results.pdf', as_attachment=True)



